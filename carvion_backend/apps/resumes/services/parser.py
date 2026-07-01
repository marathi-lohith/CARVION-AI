import io
import fitz  # PyMuPDF
import docx  # python-docx
import logging

logger = logging.getLogger("carvion.api")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract raw layout text from a PDF file in memory using PyMuPDF."""
    text_content = []
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text = page.get_text()
                if text:
                    text_content.append(text)
    except Exception as exc:
        logger.exception("PyMuPDF failed to extract text from PDF stream: %s", str(exc))
        raise ValueError("Could not extract layout text from the provided PDF file.") from exc
        
    return "\n".join(text_content)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from Word document paragraphs and tables in memory using python-docx."""
    text_content = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        
        # 1. Extract text from standard text blocks
        for paragraph in doc.paragraphs:
            cleaned_text = paragraph.text.strip()
            if cleaned_text:
                text_content.append(cleaned_text)
                
        # 2. Extract text from table grids
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cleaned_cell = cell.text.strip()
                    if cleaned_cell:
                        text_content.append(cleaned_cell)
    except Exception as exc:
        logger.exception("python-docx failed to extract text from DOCX stream: %s", str(exc))
        raise ValueError("Could not extract layout text from the provided DOCX file.") from exc
        
    return "\n".join(text_content)


def extract_resume_text(file_bytes: bytes, mime_type: str) -> str:
    """
    Directs in-memory file content parsing based on detected file MIME type.
    """
    if mime_type == "application/pdf":
        return extract_text_from_pdf(file_bytes)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file_bytes)
    else:
        logger.warning("Unsupported file MIME format block attempted: %s", mime_type)
        raise ValueError("Unsupported document format. Only PDF and DOCX uploads are accepted.")
